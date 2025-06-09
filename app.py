
from flask import Flask, request, send_file, make_response
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
import io
import os
from datetime import datetime

app = Flask(__name__)
CORS(app, resources={r"/tailor": {"origins": "chrome-extension://*"}})

@app.route('/tailor', methods=['POST', 'OPTIONS'])
def tailor_resume():
    origin = request.headers.get("Origin", "*")

    if request.method == "OPTIONS":
        response = make_response()
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response

    data = request.json
    experience = data.get("experience", [])
    skills = data.get("skills", "")

    doc = Document("base_resume.docx")

    def replace_last_n_paragraphs(section_title, new_bullets, count):
        found_index = None
        for i, para in enumerate(doc.paragraphs):
            if section_title in para.text:
                found_index = i
                break

        if found_index is None:
            print(f"❌ Section '{section_title}' not found.")
            return

        section_indices = []
        for j in range(found_index + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text.strip()
            if len(text) > 0 and text.isupper():
                break
            if text:
                section_indices.append(j)

        if len(section_indices) < count:
            print(f"⚠️ Not enough paragraphs to replace under '{section_title}'.")
            return

        for k in range(count):
            idx = section_indices[-count + k]
            clean_bullet = new_bullets[k].replace("â€¢", "").replace("•", "•").strip()
            doc.paragraphs[idx].text = clean_bullet
            for run in doc.paragraphs[idx].runs:
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"

    replace_last_n_paragraphs("iCONSULT COLLABORATIVE, SYRACUSE UNIVERSITY", experience[0:2], 2)
    replace_last_n_paragraphs("FRAPPE TECHNOLOGIES PRIVATE LIMITED", experience[2:5], 3)
    replace_last_n_paragraphs("ERNST & YOUNG", experience[5:8], 3)

    for i, para in enumerate(doc.paragraphs):
        if "Core Competencies" in para.text:
            para.text += " | " + skills
            break

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = "Rijul Chaturvedi Resume - " + datetime.now().strftime('%Y-%m-%d') + ".docx"

    response = make_response(send_file(
        output,
        as_attachment=True,
        download_name=filename
    ))
    response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)

